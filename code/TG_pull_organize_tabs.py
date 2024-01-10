import os
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
import re

folder_path = 'D:/nemo_outputs/ANALYSIS/SalishSea1500-RUN216/PLOTS/TG/year2018'

# Define the station names you want to extract
station_names = [
    'Campbell River', 'Cherry Point', 'Friday Harbour', 'Kitsilano',
    'Nanaimo Harbour', 'Patricia Bay', 'Point Atkinson', 'Port Angeles',
    'Port Renfrew', 'Port Townsend', 'Sand Heads', 'Sandy Cove', 'Victoria Harbour'
]

# Define the tidal constituents you want to extract
tidal_constituents = ['K1', 'O1', 'P1', 'Q1', 'M2', 'K2', 'N2', 'S2']

# Regular expression pattern to match the desired CSV filenames
file_pattern = re.compile(r'TG_year2018_(\w+)_RUN216.csv')

# Initialize empty dataframes for each parameter
tidal_error_df = pd.DataFrame(columns=['Station name'] + tidal_constituents)
amp_mod_obs_df = pd.DataFrame(columns=['Station name'] + tidal_constituents)
phase_mod_obs_df = pd.DataFrame(columns=['Station name'] + tidal_constituents)

# Loop over CSV files in the folder
for filename in os.listdir(folder_path):
    if filename.endswith('.csv'):
        match = file_pattern.match(filename) # match basic pattern
        tidal_constituent = match.group(1)
        if match and tidal_constituent in tidal_constituents:


            # Read the CSV file into a pandas dataframe
            file_path = os.path.join(folder_path, filename)
            df = pd.read_csv(file_path, header=0, delimiter=',')

            # Filter rows based on station names
            df = df[df['Station name'].isin(station_names)]

            # Extract and store the required data in respective dataframes
            tidal_error_df[tidal_constituent] = df['Tidal error']
            tidal_error_df['Station name'] = df['Station name']

            amp_mod_obs_df[tidal_constituent] = df['Amp (mod-obs)']
            amp_mod_obs_df['Station name'] = df['Station name']

            phase_mod_obs_df[tidal_constituent] = df['Phase (mod-obs)']
            phase_mod_obs_df['Station name'] = df['Station name']


# Create a function to add tables to a Word document
def add_table_to_doc(doc, df, table_title):
    doc.add_paragraph(table_title, style='Heading1')
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])

    # Set column names
    for i, col_name in enumerate(df.columns):
        table.cell(0, i).text = col_name

    # Fill in data
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            table.cell(i + 1, j).text = str(df.iloc[i, j])


# Create a Word document
doc = Document()
doc.add_heading('Tidal Data', level=1)

# Add tables to the document
add_table_to_doc(doc, tidal_error_df, 'Tidal error')
add_table_to_doc(doc, amp_mod_obs_df, 'Amp (mod-obs)')
add_table_to_doc(doc, phase_mod_obs_df, 'Phase (mod-obs)')

# Save the Word document
doc.save('tidal_data.docx')